"""
Generates System_Architecture_and_Gap_Analysis.docx for the Aura Exam platform.
Run: python generate_docs.py
Requires: pip install python-docx  (auto-installed if missing)
"""

import subprocess, sys

def ensure_dependency(pkg):
    try:
        __import__(pkg.replace("-", "_"))
    except ImportError:
        print(f"Installing {pkg}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

ensure_dependency("python-docx")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Colours ───────────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
DARK_BLUE  = RGBColor(0x1A, 0x23, 0x7E)
DARK_RED   = RGBColor(0xB7, 0x1C, 0x1C)
DARK_AMBER = RGBColor(0xE6, 0x5C, 0x00)
LIME_GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY       = RGBColor(0x42, 0x42, 0x42)
MID_GREY   = RGBColor(0x75, 0x75, 0x75)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def h1(text, color=DARK_BLUE):
    p = doc.add_heading(level=1)
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=18, bold=True, color=color)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(text, color=DARK_BLUE):
    p = doc.add_heading(level=2)
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True, color=color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text, color=BLACK):
    p = doc.add_heading(level=3)
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def body(text, bold=False, italic=False, color=BLACK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(4)

def bullet(text, level=0, color=BLACK):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=11, color=color)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)

def sub_bullet(text, color=MID_GREY):
    bullet(text, level=1, color=color)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(0x37, 0x47, 0x51)

def divider():
    p = doc.add_paragraph("─" * 75)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def page_break():
    doc.add_page_break()

def roadmap_item(priority, title, desc, color):
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{priority}]  ")
    r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = color
    r2 = p.add_run(title)
    r2.font.name = "Calibri"; r2.font.size = Pt(11); r2.font.bold = True; r2.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(0)
    p2 = doc.add_paragraph()
    r3 = p2.add_run(f"    {desc}")
    r3.font.name = "Calibri"; r3.font.size = Pt(10); r3.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
r = p.add_run("AURA EXAM PLATFORM")
r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = DARK_BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
r2 = p2.add_run("System Architecture, Business Rules & Gap Analysis")
r2.font.name = "Calibri"; r2.font.size = Pt(16); r2.font.color.rgb = GREY
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
r3 = p3.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}")
r3.font.name = "Calibri"; r3.font.size = Pt(11); r3.font.color.rgb = MID_GREY
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(12)

p4 = doc.add_paragraph()
r4 = p4.add_run("Confidential — Internal Use Only")
r4.font.name = "Calibri"; r4.font.size = Pt(10); r4.font.italic = True; r4.font.color.rgb = MID_GREY
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

h1("1. EXECUTIVE SUMMARY")
divider()

body(
    "Aura Exam is a full-stack online examination platform purpose-built for Malaysian tertiary "
    "classrooms. It supports the complete academic workflow — from course creation and question "
    "authoring through live examination with AI-assisted proctoring, automated scoring, appeal "
    "resolution, and grade reporting. The platform enforces academic integrity through a three-strike "
    "system, supports two assessment types (formal examinations and file-based assignments), and "
    "maintains a full audit trail of all high-stakes administrative actions."
)

h3("Technology Stack")
bullet("Full-stack: TanStack Start (React 19, SSR, file-based routing) with TypeScript 5.8")
bullet("Backend: Supabase — PostgreSQL 15, Row Level Security, Auth, Storage")
bullet("Build: Vite 7 + @lovable.dev/vite-tanstack-config")
bullet("UI: Tailwind CSS v4 (oklch design tokens), shadcn/ui, neo-brutalist design system")
bullet("Key libraries: TanStack Query, TanStack Router, Lucide React, Sonner (toasts)")

h3("Three User Roles")
bullet("Student — enrols in classes, takes exams, submits assignments, files appeals, views grades and notifications")
bullet("Lecturer — creates and manages classes, builds question banks, creates/publishes/grades exams, reviews assignments, resolves appeals")
bullet("Admin — platform oversight: user management (ban/unban), class deletion, read-only exam monitoring, audit log, PDPA compliance")

h3("Current Maturity")
body(
    "The core examination workflow is production-ready. The platform has a robust integrity "
    "enforcement pipeline, a complete appeals system, a real-time audit log, and CSV/PDF export "
    "reporting. Two admin pages (Settings, PDPA) are UI stubs with no backend. No automated "
    "test suite exists. The platform runs on Supabase Free Tier (Nano compute, South Asia / "
    "Mumbai, ap-south-1)."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("2. SYSTEM ARCHITECTURE & OVERALL SUMMARY")
divider()

h2("2.1 Data Flow & Component Mapping")

body(
    "All user interactions originate in React route components. Components call server functions "
    "(thin RPC wrappers declared with createServerFn()) which execute on the Node.js server, "
    "authenticate the caller via Supabase Auth, and query PostgreSQL. No REST endpoints or "
    "GraphQL layer exist — server functions ARE the API layer."
)
code_block(
    "Browser (React UI)\n"
    "  ↓  createServerFn() RPC call\n"
    "Node.js Server\n"
    "  ↓  supabase.auth.getUser()  →  verify identity\n"
    "  ↓  Supabase PostgreSQL       →  read / write data (RLS enforced)\n"
    "  ↓  Supabase Storage          →  class files, assignment files, proctor snapshots\n"
    "  ↓  Supabase Auth Admin API   →  user listing, audit writes (service role only)"
)

h3("Three Supabase Clients")
bullet("client.ts — browser-side, anon key, session from localStorage. Used for file uploads.")
bullet("server.ts — SSR, reads session from HTTP cookies. Used in ALL server functions.")
bullet("Admin client (audit.ts, users.ts) — service role key, bypasses RLS. Used for admin user listing and audit log writes.")

h2("2.2 Core Module Summary")

h3("src/lib/supabase/ — Server Functions (All Database Access)")
bullet("exams.ts (~1,350 lines) — exam CRUD, lifecycle, taking, scoring, grading, integrity flags, admin oversight")
bullet("classes.ts (~600 lines) — class management, enrollments, announcements, notes, assignments, assignment grading")
bullet("appeals.ts — appeal submission, deadline checks, lecturer resolution, score correction, retake flow")
bullet("audit.ts — writeAudit() helper + getAuditLog() admin server function; uses service role")
bullet("users.ts — admin user listing (Auth Admin API), ban/unban with audit logging")
bullet("proctor.ts — heartbeat pings, advisory proctoring events, signed snapshot URL generation")
bullet("funcs.ts — dashboard aggregation queries (student, lecturer, admin)")
bullet("exports.ts — CSV and print report queries for class-level reporting")
bullet("notifications.ts — push notifications, read status, mark-all-read")

h3("src/routes/_authenticated/ — Role-Scoped Pages")
bullet("student/ — dashboard, classes (detail with tabs: announcements, notes, exams, assignments), exams (list, lobby, take, result), appeals, notifications, profile")
bullet("lecturer/ — dashboard, classes (detail with 5 tabs incl. Exams), question-bank, exams (list, new, edit, results, monitor), appeals, notifications, profile")
bullet("admin/ — dashboard, users (ban/unban with modal), classes (search + delete), exams (read-only), settings (stub), audit-log (real data with category tabs), pdpa (stub)")

h3("src/components/brand/ — Custom UI Primitives")
bullet("AppShell — role-aware navigation, unread badge (polled 30s), fullscreen suppression for exam take")
bullet("ExamBuilder — multi-step exam wizard (class, schedule, questions, draft vs publish)")
bullet("QuestionEditor — MCQ, True/False, Essay creation with difficulty, points, and tags")
bullet("CameraProctor — client-side face detection, tab-switch detection, copy-paste detection")
bullet("ConfirmModal — reusable destructive-action confirmation modal (replaces all window.confirm)")
bullet("ProfilePage / NotificationsPage — shared components used by both student and lecturer")

h2("2.3 Database Schema Summary")

h3("Key Tables and Their Purpose")
bullet("profiles — user name, role (student/lecturer/admin), status (active/banned/pending)")
bullet("classes + class_enrollments — classes owned by lecturers; students enrol via a join code")
bullet("questions + exam_questions — question bank scoped to class; ordered linking to exams")
bullet("exams — full lifecycle from draft to graded; tracks status, schedule, camera requirement")
bullet("submissions — one per student per exam; tracks answers, score, flags, timing")
bullet("essay_answers — individual essay responses with manual score from lecturer")
bullet("flag_reasons — per-flag detail (type, time, label) for integrity review")
bullet("class_assignments + assignment_submissions — file-based assignments with deadlines")
bullet("appeals — one per submission; type = score | integrity; lecturer resolution stored")
bullet("notifications — user inbox; kind = exam | grade | appeal | system | integrity")
bullet("audit_log — immutable append-only log; actor_id, action, target, category, created_at")
bullet("class_notes + announcements — course materials and updates per class")

h3("Storage Buckets")
bullet("class-notes — lecture materials (PDFs, images). Currently public URL — should be private.")
bullet("class-materials — assignment attachments from lecturers and student submissions")
bullet("proctor-snapshots — face capture images scoped to {submissionId}/; served via signed URLs")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ARCHITECTURAL PROBLEMS & TECH DEBT
# ══════════════════════════════════════════════════════════════════════════════

h1("3. ARCHITECTURAL PROBLEMS & TECHNICAL DEBT")
divider()

h2("3.1 Code Quality & Bottlenecks")

h3("Pervasive any-casting bypasses all TypeScript safety")
body(
    "Every server function uses: const db = (supabase) => supabase as any. "
    "database.types.ts exists but is stale — not regenerated after schema additions. "
    "Any column name typo or wrong type will only surface as a runtime error."
)
bullet("Impact: all database queries are untyped; refactoring confidence is low", color=DARK_RED)
bullet("Fix: run supabase gen types typescript after every migration; remove the db() cast", color=LIME_GREEN)

h3("N+1 query patterns in dashboard functions")
body(
    "getLecturerDashboardData fetches classes, then exams, then loops through submission IDs to "
    "count ungraded essays across multiple round trips. The student dashboard has similar chained queries."
)
bullet("Impact: dashboards become progressively slower as data grows", color=DARK_RED)
bullet("Fix: consolidate into PostgreSQL views or Supabase RPC with aggregation", color=LIME_GREEN)

h3("No pagination on list endpoints")
body("getAllUsers, getAuditLog (limited to 500), getClassMembers, getNotifications, and question bank queries return unbounded results.")
bullet("Impact: large institutions will see slow page loads and high memory use", color=DARK_RED)
bullet("Fix: add limit/offset parameters to all list server functions", color=LIME_GREEN)

h3("Mock data file still present")
body("src/lib/mock-data.ts remains in the codebase. The auditLog export is no longer used but the file adds confusion about what is real vs stubbed.")
bullet("Fix: audit imports and delete mock-data.ts", color=LIME_GREEN)

h3("exams.ts is 1,350 lines — too large to maintain")
body("A single file covers exam creation, lifecycle, taking, scoring, grading, proctoring, and admin oversight. Difficult to navigate and test.")
bullet("Fix: split into exams-lifecycle.ts, exams-taking.ts, exams-grading.ts, exams-admin.ts", color=LIME_GREEN)

h2("3.2 Security Gaps")

h3("Admin role check missing in getAdminDashboardData")
body("getAdminDashboardData verifies the user is authenticated but does NOT check that role = 'admin'. Any logged-in user who knows the endpoint can call it.")
bullet("Impact: data leakage of platform-wide user/class/exam counts", color=DARK_RED)
bullet("Fix: add profile.role = 'admin' check at the top of the handler", color=LIME_GREEN)

h3("SUPABASE_SERVICE_ROLE_KEY not validated at startup")
body("The service role key is read lazily in createAdminClient(). A missing key causes 'Invalid API key' at request time and silently swallows all audit log writes.")
bullet("Fix: validate required env vars at server startup; fail fast with a clear error", color=LIME_GREEN)

h3("class-notes bucket serves public URLs")
body("Class materials are uploaded to a public storage bucket. Anyone with the URL can download sensitive lecture materials without being authenticated.")
bullet("Impact: course content is publicly accessible by URL", color=DARK_RED)
bullet("Fix: switch to private bucket; generate signed URLs (30 min TTL) for downloads", color=LIME_GREEN)

h3("No rate limiting on server functions")
body("recordFlag, submitExam, submitAppeal, and banUser have no rate limiting. They are vulnerable to brute-force or flooding attacks.")
bullet("Fix: implement per-user rate limiting at the middleware / Edge Function layer", color=LIME_GREEN)

h3("Client-side proctoring is bypassable")
body(
    "Tab-switch, copy-paste, and face-detection events are all client-side JavaScript. "
    "A technically capable student can disable event listeners entirely. The 3-strike system "
    "relies on client-reported events."
)
bullet("Impact: determined cheaters can bypass all integrity detection", color=DARK_RED)
bullet("Fix (partial): move hard-flag recording to a server-verified token or challenge mechanism", color=LIME_GREEN)

h2("3.3 Performance Issues")

h3("syncExamStatuses has no scheduled trigger")
body("Exam status transitions (upcoming→live→closed) rely on syncExamStatuses being called by client navigation. If no users are active, exams never transition.")
bullet("Fix: deploy a Supabase Edge Function on a 1-minute cron to call syncExamStatuses", color=LIME_GREEN)

h3("Unread notification count polled every 30 seconds")
body("AppShell fires a database query every 30 seconds per active user. This creates steady background load that scales with concurrent users.")
bullet("Fix: replace with Supabase Realtime subscription on notifications table INSERT", color=LIME_GREEN)

h3("No query caching beyond React state")
body("TanStack Query is installed but most reads go through loader() which bypasses the query cache on navigation. Repeated page visits always re-fetch from the database.")
bullet("Fix: migrate hot read paths (exams list, classes list) to useQuery with a reasonable staleTime", color=LIME_GREEN)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MISSING RULES & REQUIREMENTS GAPS
# ══════════════════════════════════════════════════════════════════════════════

h1("4. MISSING RULES & REQUIREMENTS GAPS")
divider()

h2("4.1 Stubbed / Incomplete Features")

h3("Admin Settings — no backend implementation")
body(
    "The settings page shows inputs for grace period, flag threshold, session timeout, and appeal "
    "window. The Save button does nothing. Flag threshold is hardcoded to 3 in recordFlag(); "
    "appeal window is hardcoded to 7 days in submitAppeal()."
)
bullet("Required: create a platform_settings table; wire Save button to read/write it")
bullet("Required: replace all hardcoded business constants with reads from platform_settings")

h3("PDPA / Data Deletion — no backend")
body(
    "The PDPA page shows hardcoded mock requests. Process and Decline buttons do nothing. "
    "Malaysia's PDPA 2010 requires actual data deletion or anonymisation on valid requests."
)
bullet("Required: student-initiated deletion request flow")
bullet("Required: admin processing that anonymises profile + deletes or zeros submissions")

h3("Retake flow is incomplete after integrity appeal approval")
body(
    "When an integrity appeal is approved, submission status becomes retake-approved. "
    "However, getExamForTaking only handles in-progress status — a retake-approved student "
    "cannot actually start a new attempt because no fresh submission is created."
)
bullet("Required: detect retake-approved on the lobby page; create a new submission on entry")

h2("4.2 Edge Case Discrepancies")

h3("Exam can be published with zero questions")
body("updateExam accepts an empty question_ids array with status = 'upcoming'. Students would see a blank exam page.")
bullet("Fix: server-side enforce at least 1 question before allowing upcoming status")

h3("Appeal deadline is NOT enforced server-side")
body("submitAppeal only shows/hides the appeal UI client-side. The server function does not check whether the 7-day window has expired. An API call can file a late appeal.")
bullet("Fix: add deadline check inside the submitAppeal handler")

h3("Assignment max score is not enforced server-side")
body("The UI caps the grade input at max_score, but reviewAssignmentSubmission accepts any numeric grade. A lecturer could award more than the maximum.")
bullet("Fix: fetch assignment.max_score in the handler and reject over-limit grades")

h3("Notification list is not paginated")
body("getNotifications returns all notifications for a user with no limit. Long-term users of active courses could accumulate hundreds of records, all loaded on every inbox visit.")

h3("essay_answers.score — no upper bound validation server-side")
body("gradeEssayAnswer does not check the question's points value. A typo (e.g. 100 instead of 10) could award 10x the intended score.")
bullet("Fix: fetch question.points and reject scores exceeding it")

h2("4.3 Business Constants — Hardcoded, Not Configurable")
body("The following values exist only in code and cannot be changed without a deployment:")
bullet("Flag threshold: 3 strikes → auto-submit  (recordFlag, exams.ts)")
bullet("Appeal window: 7 days  (submitAppeal, appeals.ts)")
bullet("Heartbeat interval: 30 seconds  (exam take page)")
bullet("Notification poll interval: 30 seconds  (AppShell)")
bullet("Proctor snapshot signed URL TTL: 30 minutes  (getProctorSnapshots)")
bullet("Assignment resubmission: unlimited before review  (no max attempt cap)")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — AI INTEGRATION OPPORTUNITIES
# ══════════════════════════════════════════════════════════════════════════════

h1("5. FUTURE-PROOFING & AI INTEGRATION OPPORTUNITIES")
divider()

h2("5.1 Immediate AI Wins (Low Effort, High Value)")

h3("1. AI Essay Grading Assistance")
body(
    "The platform has a manual essay grading UI. An LLM can pre-score each essay, "
    "supply a rationale, and flag unusual answers. Lecturers confirm or override. "
    "This reduces grading time by 60-80% for large cohorts."
)
bullet("Integration: inside gradeEssayAnswer — call Claude after submission closes, store suggested_score + rationale")
bullet("Model: claude-sonnet-4-6 (quality) or claude-haiku-4-5 (cost-optimised)")
bullet("Prompt: question text + marking rubric + student essay + class context")
bullet("UX: show AI suggestion with one-click accept in the grading panel")

h3("2. Smart Question Tagging & Difficulty Prediction")
body("When a lecturer creates a question, auto-suggest Bloom's Taxonomy level, difficulty, and topic tags based on question text.")
bullet("Integration: after createQuestion / updateQuestion")
bullet("Model: claude-haiku-4-5 (fast and cheap for classification)")
bullet("Benefit: makes the question bank fully searchable; enables AI-balanced exam generation")

h3("3. Automated Post-Exam Performance Reports")
body("After all essays are graded, generate a natural-language summary per student (strengths, weak topics, recommended revision areas) and a class-level analysis for the lecturer.")
bullet("Integration: background job triggered when submission.status = 'graded'")
bullet("Output: stored as notification + downloadable PDF")

h3("4. Appeal Reason Classification & Triage")
body("Appeal reasons are free text. An LLM can classify them (legitimate technical issue vs vague claim vs strong dispute) and sort the lecturer queue by case strength.")
bullet("Integration: submitAppeal — classify immediately, store as appeals.triage_label")

h2("5.2 Medium-Term AI Enhancements")

h3("5. Academic Integrity Content Analysis")
body(
    "Check essay answers for AI-generated content patterns and semantic similarity "
    "across the class cohort. Flags cluster copying or suspiciously uniform phrasing."
)
bullet("Stack: pgvector (Supabase extension) + text embeddings + cosine similarity queries")
bullet("Threshold: flag submissions with > 85% similarity to classmate answers")

h3("6. Semantic Question Bank Search")
body("Enable 'find questions similar to this topic' and auto-detect near-duplicate questions using vector embeddings rather than keyword search.")
bullet("Stack: pgvector + embedding generation on question create/update")

h2("5.3 Architecture Readiness for AI")

h3("What already supports AI integration")
bullet("Server functions provide a secure, server-only execution context — no API key exposure to browser")
bullet("Supabase RLS means AI-generated data stays correctly scoped without extra auth logic")
bullet("essay_answers.score is already nullable — AI can write suggested scores without schema changes")
bullet("The audit log can naturally extend to record AI-assisted actions")

h3("What needs to be added")
bullet("@anthropic-ai/sdk package — add to package.json; implement with prompt caching for repeated system prompts")
bullet("pgvector extension — enable on Supabase for semantic search (free on Pro, $0 on paid plan)")
bullet("Background job runner — Supabase Edge Functions or a queue (Inngest, Upstash) for async AI tasks")
bullet("New schema columns: questions.ai_difficulty, questions.ai_tags, essay_answers.ai_score, essay_answers.ai_rationale, submissions.ai_integrity_score, appeals.ai_triage")
bullet("Cost controls: per-request token limits; log AI usage in audit_log")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — ACTIONABLE ROADMAP
# ══════════════════════════════════════════════════════════════════════════════

h1("6. ACTIONABLE ROADMAP")
divider()

body("Priority levels:  HIGH = blocks correctness or compliance  |  MEDIUM = reliability / UX  |  LOW = quality / future-proofing", italic=True)
doc.add_paragraph()

h2("PHASE 1 — Critical Fixes (Do This Week)")

roadmap_item("HIGH", "Apply pending SQL migrations",
    "Run the two SQL statements in Supabase SQL Editor: add audit_log.category column and update class_enrollments RLS to allow admin reads. Without these, the audit log is broken and student counts on admin/classes show zero.",
    DARK_RED)

roadmap_item("HIGH", "Fix admin role check in getAdminDashboardData",
    "Add a profile.role = 'admin' check. Any authenticated user can currently call this endpoint.",
    DARK_RED)

roadmap_item("HIGH", "Validate env vars at startup",
    "Check SUPABASE_SERVICE_ROLE_KEY exists at server startup and throw a clear error if missing. Prevents silent audit log failures.",
    DARK_RED)

roadmap_item("HIGH", "Enforce appeal deadline server-side",
    "Add a 7-day window check inside submitAppeal handler. The client-side guard is not sufficient.",
    DARK_RED)

roadmap_item("HIGH", "Enforce minimum 1 question before exam publish",
    "Add server-side validation in updateExam: reject status = 'upcoming' if question_ids is empty.",
    DARK_RED)

h2("PHASE 2 — Reliability & Security (Next Sprint)")

roadmap_item("MEDIUM", "Deploy syncExamStatuses cron job",
    "Create a Supabase Edge Function on a 1-minute cron schedule. Exam status transitions currently only happen on client navigation.",
    DARK_AMBER)

roadmap_item("MEDIUM", "Switch class-notes bucket to private",
    "Change class-notes to private storage. Generate 30-minute signed URLs for downloads. Prevents public URL access to sensitive lecture materials.",
    DARK_AMBER)

roadmap_item("MEDIUM", "Fix retake flow for integrity-approved appeals",
    "On the student lobby page, detect retake-approved status and create a fresh submission instead of resuming the flagged one.",
    DARK_AMBER)

roadmap_item("MEDIUM", "Add max score validation in gradeEssayAnswer",
    "Fetch question.points and reject scores that exceed it. Also enforce assignment.max_score in reviewAssignmentSubmission.",
    DARK_AMBER)

roadmap_item("MEDIUM", "Implement Admin Settings backend",
    "Create a platform_settings table. Wire the save button to a server function. Replace hardcoded flag threshold (3) and appeal window (7 days) with reads from this table.",
    DARK_AMBER)

roadmap_item("MEDIUM", "Implement PDPA data deletion flow",
    "Replace the mock PDPA page with a real student-initiated deletion request and admin processing workflow (anonymise profile + zero submissions).",
    DARK_AMBER)

roadmap_item("MEDIUM", "Hide submission UI for reviewed assignments",
    "Detect reviewed status in the student class view and replace the upload form with a read-only grade/feedback display.",
    DARK_AMBER)

h2("PHASE 3 — Code Quality (Following Sprint)")

roadmap_item("MEDIUM", "Regenerate database.types.ts and remove db() cast",
    "Run supabase gen types typescript. Replace the db() = supabase as any pattern. This restores TypeScript type safety across all database queries.",
    LIME_GREEN)

roadmap_item("MEDIUM", "Delete src/lib/mock-data.ts",
    "Audit all imports and confirm no pages still use mock data. Remove the file. The auditLog export is already unused.",
    LIME_GREEN)

roadmap_item("LOW", "Add pagination to list endpoints",
    "Add limit + offset to getAllUsers, getAuditLog, getNotifications, getClassMembers, getAssignmentSubmissions.",
    LIME_GREEN)

roadmap_item("LOW", "Replace notification polling with Supabase Realtime",
    "Subscribe to notifications INSERT in AppShell. Remove the 30-second polling timer.",
    LIME_GREEN)

roadmap_item("LOW", "Split exams.ts into domain modules",
    "Separate into exams-lifecycle.ts, exams-taking.ts, exams-grading.ts, exams-admin.ts for maintainability.",
    LIME_GREEN)

h2("PHASE 4 — AI Features (Future)")

roadmap_item("MEDIUM", "AI essay grading assistance (Phase 1)",
    "Add @anthropic-ai/sdk. On essay closed, call Claude claude-sonnet-4-6 to generate a suggested score + rationale. Show in the grading panel with one-click accept.",
    DARK_BLUE)

roadmap_item("MEDIUM", "Smart question tagging",
    "Call Claude claude-haiku-4-5 on question create/update to auto-suggest difficulty and topic tags. Store in questions table.",
    DARK_BLUE)

roadmap_item("LOW", "Academic integrity content analysis",
    "Add pgvector, generate essay embeddings, compute intra-class similarity, flag outliers for lecturer review.",
    DARK_BLUE)

roadmap_item("LOW", "Post-exam performance reports",
    "Generate natural-language summaries per student after grading. Deliver as notifications + downloadable PDF.",
    DARK_BLUE)

roadmap_item("LOW", "Appeal triage AI",
    "Classify appeal reasons using Claude on submitAppeal. Store triage_label. Sort lecturer appeal queue by urgency.",
    DARK_BLUE)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A — COMPLETE BUSINESS RULES
# ══════════════════════════════════════════════════════════════════════════════

h1("APPENDIX A — COMPLETE BUSINESS RULES REFERENCE")
divider()

h2("A.1 Exam Lifecycle")
bullet("Created in draft status — invisible to students")
bullet("Publishing (draft → upcoming) requires: title, class, start_time, end_time, duration > 0, ≥1 question")
bullet("Once published (upcoming): only title and require_camera may be changed")
bullet("Unpublishing (upcoming → draft): only allowed if no submissions AND start_time has not passed")
bullet("live: starts automatically when the server clock passes start_time")
bullet("closed: starts automatically when end_time passes")
bullet("graded: when all essay_answers have a score (MCQ/TF-only exams go straight to graded on close)")
bullet("live, closed, graded: cannot be edited or deleted")

h2("A.2 Submission & Scoring")
bullet("One active submission per student per exam")
bullet("auto_score = MCQ + TF points; calculated at submit time; immutable thereafter")
bullet("score = auto_score + sum of graded essay scores; updated each time an essay is scored")
bullet("MCQ: correct = full points; incorrect / unanswered = 0. No partial credit.")
bullet("True/False: correct = full points; incorrect = 0")
bullet("Essay: score = null until manually graded; submission stays 'submitted' until ALL essays have scores")
bullet("Timer is seeded from started_at + duration, bounded by end_time (server-calculated on every loader call)")
bullet("Answers and question index persisted to sessionStorage keyed by submissionId; retake always starts fresh")

h2("A.3 Integrity & Proctoring")
bullet("Hard events (tab-switch, copy-paste): increment flags counter; recorded in flag_reasons")
bullet("Advisory events (face-not-visible, multiple-faces, camera-lost): recorded in flag_reasons; DO NOT increment flags")
bullet("flags = 3: auto-submit with status = 'flagged', score = 0; student notified with 7-day appeal window")
bullet("require_camera flag per exam: student must grant camera permission before entering")
bullet("Face snapshots stored in proctor-snapshots bucket under {submissionId}/; accessible to student and class lecturer")
bullet("Client-side detection is acknowledged as bypassable — it is deterrence + evidence, not enforcement")

h2("A.4 Appeals")
bullet("One appeal per submission; type = score (dispute grading) or integrity (challenge auto-submit)")
bullet("Appeal window: 7 days from submission date — late appeals blocked (currently client-side only; server enforcement is a gap)")
bullet("Starts pending; only the exam's class lecturer may resolve it")
bullet("Score appeal approved: lecturer must supply a corrected score; submissions.score updated immediately")
bullet("Integrity appeal approved: submission → retake-approved; student may retake (retake flow is incomplete — see gap)")
bullet("Rejected: lecturer reply recorded; original score/status stands")
bullet("Both outcomes trigger a student notification")

h2("A.5 Assignments")
bullet("Has start_at (visible from) and end_at (submission deadline); both required")
bullet("Server enforces deadline: submissions after end_at are rejected regardless of client")
bullet("Student may resubmit any number of times before deadline; each resubmission replaces the file and resets status to 'submitted'")
bullet("After lecturer marks 'reviewed': submission is permanently locked; resubmission rejected with error")
bullet("Grading is only available after end_at has passed")
bullet("Optional max_score; enforced in UI but NOT server-side (gap)")

h2("A.6 Audit Log")
bullet("Written using service role key; never throws (logging failures must not break primary operations)")
bullet("Immutable append-only: no UPDATE or DELETE policies on audit_log")
bullet("Readable only by role = admin")
bullet("Logged events: user ban/unban, class deletion, exam publish/unpublish/delete, all essays graded, exam auto-submitted by integrity, appeal approved/rejected")
bullet("NOT logged: routine reads, draft saves, enrollments, announcements, logins")

h2("A.7 User Management")
bullet("Registration defaults to role = student")
bullet("Lecturer accounts require admin approval (status: pending → active)")
bullet("Banned users (status = banned) cannot log in; application checks profile.status on every authenticated request")
bullet("Admins cannot ban other admins")
bullet("No user deletion; accounts are banned to preserve submission history")

h2("A.8 Notifications")
bullet("Sent on: exam published, exam graded, exam auto-submitted, appeal approved, appeal rejected, assignment submitted (to lecturer), assignment reviewed (to student), new announcement, new note")
bullet("kind: exam | grade | appeal | system | integrity")
bullet("Unread count polled every 30 seconds by AppShell")
bullet("Students and lecturers can mark individual or all notifications as read")
bullet("Not paginated (potential performance issue with long-term users)")

h2("A.9 Classes & Enrollments")
bullet("Class created by lecturer with a name, a randomly generated join code, and a colour")
bullet("Students join using the 6-character code; duplicate enrollment rejected")
bullet("Lecturer can remove students from a class (unenrolment)")
bullet("Admin can delete a class; deletion cascades to all related data (exams, submissions, assignments)")
bullet("Admin read of class_enrollments requires the updated RLS policy (see pending migrations)")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B — ENVIRONMENT & DEPLOYMENT CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════

h1("APPENDIX B — ENVIRONMENT & DEPLOYMENT CHECKLIST")
divider()

h2("Required Environment Variables")
code_block(
    "VITE_SUPABASE_URL=https://<project-ref>.supabase.co\n"
    "VITE_SUPABASE_ANON_KEY=<anon key>          # public — safe to expose to browser\n"
    "SUPABASE_SERVICE_ROLE_KEY=<service_role>   # secret — server-only, never expose to browser\n"
    "DATABASE_URL=<transaction pooler URL>      # only needed by scripts/migrate-*.ts"
)

h2("Storage Buckets")
bullet("class-notes — currently public (should be private — see security gap)")
bullet("class-materials — private")
bullet("proctor-snapshots — private")

h2("Pending SQL Migrations (run in Supabase SQL Editor)")
code_block(
    "-- 1. audit_log category column\n"
    "ALTER TABLE audit_log\n"
    "  ADD COLUMN IF NOT EXISTS category text NOT NULL DEFAULT 'general'\n"
    "  CHECK (category IN ('user_management','exam','integrity','appeal','class','general'));\n\n"
    "-- 2. admin reads class_enrollments (fixes zero student count on admin/classes)\n"
    "DROP POLICY IF EXISTS \"enrollments: read\" ON class_enrollments;\n"
    "CREATE POLICY \"enrollments: read\" ON class_enrollments FOR SELECT\n"
    "  USING (\n"
    "    auth.uid() = student_id\n"
    "    OR EXISTS (SELECT 1 FROM classes WHERE id = class_id AND lecturer_id = auth.uid())\n"
    "    OR EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')\n"
    "  );"
)

h2("First-Run Setup Order")
bullet("1. Copy .env.example → .env and fill in all four variables")
bullet("2. Run both SQL migrations in Supabase Dashboard → SQL Editor")
bullet("3. npm run create:buckets")
bullet("4. npm run seed:admin")
bullet("5. npm run seed:lecturer  (optional: test lecturer account)")
bullet("6. npm run dev")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

output_path = "System_Architecture_and_Gap_Analysis.docx"
doc.save(output_path)
print(f"\n  Document saved: {output_path}")
print(f"  Sections: Executive Summary, Architecture, Tech Debt, Gaps, AI Opportunities, Roadmap")
print(f"  Appendices: Complete Business Rules (9 sections), Deployment Checklist")
