# -*- coding: utf-8 -*-
"""Generate Aura Exam - System Analysis & Recommendations .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

INK = RGBColor(0x1A, 0x1A, 0x1A)
LIME = RGBColor(0x6B, 0x8E, 0x00)
VIOLET = RGBColor(0x5B, 0x3F, 0xCC)
PINK = RGBColor(0xC0, 0x2A, 0x6B)
GREY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)

for lvl, color, size in [(1, VIOLET, 17), (2, INK, 13.5), (3, PINK, 11.5)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.color.rgb = color
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text="", size=10.5, bold=False, italic=False, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        if color:
            r.font.color.rgb = color
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def table(headers, rows, widths=None, header_fill="5B3FCC"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_fill)
        pp = hdr[i].paragraphs[0]
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.space_before = Pt(2)
        rr = pp.add_run(h)
        rr.bold = True
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "F2F0FB")
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.space_after = Pt(2)
            pp.paragraph_format.space_before = Pt(2)
            rr = pp.add_run(str(val))
            rr.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    pPr.append(pbdr)


# COVER
doc.add_paragraph().paragraph_format.space_after = Pt(60)
para("AURA EXAM", size=40, bold=True, color=VIOLET, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("System Analysis, Architecture Review & Recommendations", size=15, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Online Examination Platform for Malaysian Classrooms", size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("Prepared for: Project Owner", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Date: " + datetime.date.today().strftime("%d %B %Y"), size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Stack: TanStack Start - React 19 - Supabase - Tailwind v4", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# CONTENTS
doc.add_heading("Contents", level=1)
for item in [
    "1. Executive Summary",
    "2. Project Requirements & Design",
    "3. System Architecture",
    "4. Data Model",
    "5. Core User Flows",
    "6. Feature-by-Feature Assessment",
    "7. What's Working Well",
    "8. Critical Issues & What You're Missing",
    "9. Recommended Best Flows",
    "10. Improvement Roadmap (Prioritised)",
    "11. Additional Suggestions",
    "12. Conclusion",
]:
    para(item, size=11, space_after=4)
doc.add_page_break()

# 1. EXECUTIVE SUMMARY
doc.add_heading("1. Executive Summary", level=1)
para(
    "Aura Exam is a role-based online examination platform with three actors - student, lecturer, "
    "and admin - built on TanStack Start (SSR React) with Supabase for authentication, PostgreSQL, "
    "row-level security (RLS), and file storage. The product covers the full assessment lifecycle: "
    "class management, a reusable question bank, exam building and scheduling, proctored exam-taking "
    "with integrity flagging and webcam face-detection, automated MCQ/TF grading, manual essay grading, "
    "a structured appeals workflow, assignments, notifications, data exports, and admin oversight."
)
para(
    "The implementation is broad and genuinely functional - most screens are wired to real server "
    "functions rather than mocks, and the grading, appeals, and notification logic is server-authoritative "
    "in the right places. The codebase is clean and consistent in its patterns."
)
para("This review found the system to be strong on feature coverage but exposed to several "
     "high-impact correctness and integrity risks that should be addressed before any real exam runs. "
     "The most urgent are summarised below and detailed in Section 8.", bold=True)
para("Top 3 things to fix first:", bold=True, color=RED, space_after=3)
bullet("Correct answers for MCQ/TF questions are sent to the browser during an exam - a student can read them from the network response.", bold_lead="ANSWER LEAK - ", color=RED)
bullet("The exam countdown is purely client-side and resets to full duration on page refresh; there is no server-enforced time limit.", bold_lead="TIMER NOT ENFORCED - ", color=RED)
bullet("The generated database types file is out of sync with the live schema, hiding real columns (auto_score, read_at, link, deadline_at) and statuses.", bold_lead="SCHEMA DRIFT - ", color=RED)
hrule()

# 2. REQUIREMENTS & DESIGN
doc.add_heading("2. Project Requirements & Design", level=1)
para("This section restates the product requirements and design intent, reconstructed from the project "
     "plan and the implemented behaviour, so the document is self-contained.")

doc.add_heading("2.1 Vision & Design Direction", level=2)
para("A calm, fair, and kind online examination experience. The visual language is intentionally playful "
     "and neo-brutalist to reduce exam anxiety:")
bullet("Warm cream background (#fcfbf6) with off-white / pale-lilac surfaces.", bold_lead="Palette: ")
bullet("Neon lime (primary CTA), hot pink, electric purple, sky blue as role accents - student = sky, lecturer = violet, admin = lime.", bold_lead="Accents: ")
bullet("Chunky rounded display headings; mono font for timers and IDs.", bold_lead="Type: ")
bullet("Pill buttons with 2px borders + offset shadow, rounded cards, scattered SVG doodles.", bold_lead="Shapes: ")
para("Design tokens live as oklch semantic variables in src/styles.css.", italic=True, color=GREY)

doc.add_heading("2.2 Roles & Permissions", level=2)
table(
    ["Role", "Primary capabilities"],
    [
        ["Student", "Join classes, view notes/announcements, take exams (proctored), view results, submit assignments, file appeals, read notifications, edit profile."],
        ["Lecturer", "Manage classes & members, build a question bank, create/schedule exams, monitor live exams, auto + manual grading (essays, score override), set assignment points, review appeals, export class data."],
        ["Admin", "Platform oversight: manage users (ban/unban, approve lecturers), global class/exam views, platform settings (grace period, flag threshold, session timeout), audit log, PDPA data-deletion requests."],
    ],
    widths=[1.1, 5.4],
)

doc.add_heading("2.3 Functional Requirements", level=2)
table(
    ["#", "Requirement", "Status"],
    [
        ["FR-1", "Email/password auth with role-based home routing", "Implemented"],
        ["FR-2", "Class creation, enrolment, announcements & notes", "Implemented"],
        ["FR-3", "Question bank (MCQ / True-False / Essay) with difficulty, tags, points, versioning", "Implemented (versioning unused)"],
        ["FR-4", "Exam builder: select questions, schedule window, duration, camera requirement", "Implemented"],
        ["FR-5", "Exam lifecycle: draft -> upcoming -> live -> graded", "Implemented"],
        ["FR-6", "Proctored exam runtime: timer, question nav, flag-for-review, webcam face check", "Implemented (client-side)"],
        ["FR-7", "Integrity monitoring: tab-switch / copy / paste / right-click detection, auto-submit at 3 flags", "Implemented (client-side only)"],
        ["FR-8", "Automated grading of MCQ/TF; manual essay grading; score override", "Implemented (server-side)"],
        ["FR-9", "Appeals: score dispute & integrity appeal, 7-day window, lecturer review queue", "Implemented"],
        ["FR-10", "Assignments: upload brief, student file submission, grading with max score & feedback", "Implemented"],
        ["FR-11", "In-app notifications across exam/grade/appeal events", "Implemented"],
        ["FR-12", "Data export: CSV + printable reports of class/student/exam scores", "Implemented"],
        ["FR-13", "Admin: user management, settings, audit log, PDPA requests", "Partially - verify audit-log writes"],
    ],
    widths=[0.5, 4.8, 1.7],
)

doc.add_heading("2.4 Non-Functional Requirements", level=2)
bullet("SSR for fast first paint and SEO on marketing pages; each public route sets its own head metadata.", bold_lead="Performance: ")
bullet("RLS-backed data isolation; auth enforced at the _authenticated layout via beforeLoad redirect.", bold_lead="Security: ")
bullet("Webcam video and MediaPipe face detection run locally in the browser - nothing recorded or uploaded (privacy by design).", bold_lead="Privacy: ")
bullet("PDPA data-deletion workflow present for Malaysian data-protection compliance.", bold_lead="Compliance: ")
bullet("Exam runtime is desktop-only by design (lockdown UX).", bold_lead="Platform: ")
hrule()

# 3. ARCHITECTURE
doc.add_heading("3. System Architecture", level=1)
doc.add_heading("3.1 Technology Stack", level=2)
table(
    ["Layer", "Technology"],
    [
        ["Framework", "TanStack Start (SSR) + TanStack Router (file-based) + TanStack Query"],
        ["UI", "React 19, Tailwind CSS v4, shadcn/ui, custom 'brand' primitives, lucide icons, sonner toasts"],
        ["Backend", "Supabase: Auth, PostgreSQL, RLS, Storage. Server functions via createServerFn"],
        ["Direct DB", "node-postgres (pg) Pool for admin/privileged queries that bypass RLS"],
        ["Proctoring", "MediaPipe Tasks Vision (FaceDetector) loaded from CDN, GPU delegate, local inference"],
        ["Build", "Vite via @lovable.dev/vite-tanstack-config; Cloudflare build target"],
        ["Validation", "Zod; react-hook-form for forms"],
    ],
    widths=[1.3, 5.2],
)

doc.add_heading("3.2 Request & Data Flow", level=2)
para("All data access flows through server functions in src/lib/supabase/*.ts. The canonical pattern is:")
code('export const fn = createServerFn({ method: "GET" | "POST" })\n  .inputValidator((data) => data)\n  .handler(async ({ data }) => {\n    const supabase = createClient();           // SSR client, reads session from cookies\n    const { data: { user } } = await supabase.auth.getUser();\n    if (!user) throw new Error("Unauthorized"); // authn gate\n    ...                                          // RLS enforces authz\n  });')
para("Two Supabase clients exist: a browser client (client.ts) and an SSR client (server.ts) that "
     "reads/writes the session via cookies. middleware.ts refreshes the session. Privileged operations "
     "(admin user list, profile creation) use a raw pg Pool that bypasses RLS.")

doc.add_heading("3.3 Routing", level=2)
para("File-based routes under src/routes. A pathless _authenticated layout guards every signed-in route "
     "and wraps content in the role-aware AppShell. Route tree is auto-generated (routeTree.gen.ts).")
bullet("dashboard, classes, exams (lobby/take/submit/result), appeals, notifications, profile", bold_lead="Student: ")
bullet("dashboard, classes, question-bank, exams (new/edit/monitor/results), appeals, notifications, profile", bold_lead="Lecturer: ")
bullet("dashboard, users, classes, exams, settings, audit-log, pdpa", bold_lead="Admin: ")
hrule()

# 4. DATA MODEL
doc.add_heading("4. Data Model", level=1)
para("Core PostgreSQL tables (Supabase). Relationships shown where they drive behaviour.")
table(
    ["Table", "Purpose & key columns"],
    [
        ["profiles", "User profile mirror of auth.users - id, name, role, status (active/banned/pending)"],
        ["classes", "code, name, lecturer_id, color"],
        ["class_enrollments", "(class_id, student_id) join - enrolment"],
        ["questions", "type (MCQ/TF/ESSAY), text, difficulty, points, tags, version, meta (options/correct/model_answer/rubric)"],
        ["exams", "title, class_id, start_time, end_time, duration, status, questions_count, require_camera"],
        ["exam_questions", "(exam_id, question_id, order_index) - ordered link"],
        ["submissions", "status (in-progress/submitted/graded/flagged/retake-approved), score, auto_score, total, flags, appeal_required, submitted_at"],
        ["essay_answers", "per-question answers; score NULL until graded; graded_by, graded_at"],
        ["flag_reasons", "integrity events per submission - time, type, label"],
        ["appeals", "type (score/integrity), status, reason, lecturer_reply, decided_at, deadline_at"],
        ["class_assignments / assignment_submissions", "assignment brief + student uploads with grade & feedback"],
        ["announcements / class_notes", "class feed + file attachments"],
        ["notifications", "type, title, body, link, read_at"],
        ["audit_log", "actor_id, action, target"],
    ],
    widths=[1.8, 4.7],
)
para("Scoring model (important):", bold=True, space_before=4)
bullet("auto_score = immutable sum of MCQ + TF points, computed server-side at submit.")
bullet("score = auto_score + sum of graded essay scores; submission promoted to 'graded' once all essays scored.")
bullet("Exam totals are computed live from question points, not from the (potentially stale) submissions.total column.")
hrule()

# 5. CORE FLOWS
doc.add_heading("5. Core User Flows", level=1)
doc.add_heading("5.1 Authentication", level=2)
para("Sign-up stores name + role in user_metadata and creates a profiles row via a pg-pool server "
     "function (bypassing RLS). Sign-in maps the Supabase user to an AuthUser and back-fills a profile "
     "row for legacy accounts. ROLE_HOME routes each role to its dashboard.")

doc.add_heading("5.2 Student Exam Lifecycle", level=2)
numbered("Student opens the exam lobby - sees schedule, duration, and a webcam + face-detection check (if camera required).", bold_lead="Lobby: ")
numbered("startExam creates (or resets, for an approved retake) an in-progress submission and snapshots total points.", bold_lead="Start: ")
numbered("Client renders questions with a countdown, per-question flag-for-review, and integrity listeners. Each violation calls recordFlag; the 3rd flag auto-submits as 'flagged'.", bold_lead="Take: ")
numbered("submitExam grades MCQ/TF server-side, stores essay answers as ungraded, and sets status 'submitted' with auto_score.", bold_lead="Submit: ")
numbered("Student sees score, per-question review, flag history, and retake/appeal CTAs.", bold_lead="Result: ")

doc.add_heading("5.3 Lecturer Grading", level=2)
para("The results screen lists every submission with auto_score, flags, and essay answers (plus model "
     "answer/rubric). gradeEssayAnswer saves a per-answer score, recomputes score = auto_score + essays, "
     "and once all essays are graded promotes the submission to 'graded' and notifies the student.")

doc.add_heading("5.4 Appeals", level=2)
bullet("Score dispute or integrity appeal, allowed only within 7 days of submission; integrity appeals require a flagged submission; duplicate types are blocked.", bold_lead="Filing: ")
bullet("Paginated queue ordered pending-first, integrity-before-score, oldest-first; lecturer sees the disputed answers.", bold_lead="Queue: ")
bullet("Approve a score appeal writes the corrected score; approve an integrity appeal sets 'retake-approved' so the student can restart. Student is notified either way.", bold_lead="Resolve: ")

doc.add_heading("5.5 Notifications", level=2)
para("Server functions push in-app notifications on key events (exam published, exam flagged, grade "
     "published, appeal submitted/approved/rejected). The AppShell polls the unread count every 30s.")
hrule()

# 6. FEATURE ASSESSMENT
doc.add_heading("6. Feature-by-Feature Assessment", level=1)
table(
    ["Feature", "Maturity", "Notes"],
    [
        ["Auth & roles", "Solid", "Clean; consider email verification gating & password policy"],
        ["Classes & enrolment", "Solid", "Works; enrolment by code is the natural next UX"],
        ["Question bank", "Good", "version column exists but no version history UI/logic"],
        ["Exam builder", "Good", "Scheduling present; no server validation of start<end or window vs duration"],
        ["Exam runtime", "At risk", "Answer leak + client-only timer + bypassable anti-cheat"],
        ["Grading", "Solid", "Server-authoritative; good separation of auto vs essay score"],
        ["Appeals", "Solid", "Well-modelled state machine and notifications"],
        ["Assignments", "Good", "File upload + grading; verify storage bucket RLS"],
        ["Notifications", "Good", "Polling works; realtime would be better"],
        ["Exports", "Good", "CSV + print; useful slip-style reports"],
        ["Admin", "Partial", "User mgmt via pg pool; confirm audit-log is actually written"],
        ["Proctoring (camera)", "Advisory", "Local face detection never reaches server; informational only"],
    ],
    widths=[1.6, 1.0, 3.9],
)
hrule()

# 7. WHAT'S WORKING WELL
doc.add_heading("7. What's Working Well", level=1)
bullet("Grading is server-side and tamper-resistant; the auto_score / score split is a clean, correct design.", bold_lead="Authoritative scoring. ")
bullet("Consistent createServerFn pattern with an authn gate at the top of every handler - easy to read and extend.", bold_lead="Uniform server-function pattern. ")
bullet("Appeals are a proper state machine with deadlines, duplicate prevention, ownership checks, and student notifications.", bold_lead="Well-modelled appeals. ")
bullet("Exam totals are recomputed from question points instead of trusting a stored total - avoids a whole class of stale-data bugs.", bold_lead="Defensive totals. ")
bullet("Webcam inference runs locally with nothing uploaded - a privacy-respecting default and a good PDPA story.", bold_lead="Privacy-first proctoring. ")
bullet("Feature breadth is high and screens are wired to real data, not mocks.", bold_lead="Real, broad coverage. ")
hrule()

# 8. CRITICAL ISSUES
doc.add_heading("8. Critical Issues & What You're Missing", level=1)
para("Ordered by severity. P0 = fix before any real exam; P1 = fix soon; P2 = hardening.", italic=True, color=GREY)

doc.add_heading("P0-1  Correct answers leak to the browser during an exam", level=3)
para("getExamForTaking selects each question's full meta and returns it to the client. For MCQ/TF, meta "
     "contains the correct answer (correct index / boolean). A student can open the network tab and read "
     "every answer.")
para("Fix: strip answer keys server-side before sending the exam to the student - return only what the UI "
     "needs to render (option text, no 'correct'). Keep the key server-side for grading.", color=LIME)

doc.add_heading("P0-2  Exam timer is not enforced", level=3)
para("The countdown initialises to exam.duration * 60 on the client every time the component mounts, so a "
     "refresh restarts the full timer. There is no server check of elapsed time or end_time at submit.")
para("Fix: persist a server started_at on the submission; compute remaining time from server timestamps; "
     "reject or auto-grade submissions past the allowed window at submit time.", color=LIME)

doc.add_heading("P0-3  Anti-cheat is client-side and bypassable", level=3)
para("Tab-switch/copy/paste/right-click detection lives entirely in browser event listeners and can be "
     "disabled (devtools, JS off, second device). Camera face-detection results never reach the server, "
     "so a 'no face' or 'multiple faces' state has no consequence.")
para("Fix: treat client flags as advisory; add a server heartbeat during the exam, optionally upload "
     "periodic face snapshots for review, and record face-missing events server-side. Be explicit in the "
     "product that browser-based proctoring is deterrence, not prevention.", color=LIME)

doc.add_heading("P1-1  Database types are stale (schema drift)", level=3)
para("src/lib/database.types.ts disagrees with the live schema in several places: submissions is missing "
     "auto_score and the retake-approved status; notifications is typed with kind/read but the code uses "
     "type/read_at/link; appeals is typed with exam_title/question_ref but the code uses "
     "exam_id/lecturer_reply/decided_at/deadline_at. The handlers cast db(supabase) to any, so these "
     "mismatches compile silently.")
para("Fix: regenerate types from Supabase (supabase gen types) and drop the 'as any' casts so the compiler "
     "catches column mistakes.", color=LIME)

doc.add_heading("P1-2  Authorisation depends on RLS without in-code defence in depth", level=3)
para("Some read handlers (e.g. getExam, getLecturerExamResults) verify only that a user is logged in and "
     "rely on RLS for row scoping. If an RLS policy is ever misconfigured, data leaks. Appeals handlers do "
     "this well (explicit ownership checks) - apply the same pattern everywhere.")
para("Fix: add explicit ownership/enrolment checks in handlers that return sensitive rows, in addition to RLS.", color=LIME)

doc.add_heading("P1-3  Leftover debug endpoint", level=3)
para("getTestProfiles returns up to 10 profile rows and appears to be debug scaffolding. Remove it (or "
     "gate behind admin) before production.")

doc.add_heading("P2  Other gaps", level=3)
bullet("No automated tests of any kind - grading, appeal windows, and flag thresholds are exactly the logic that deserves unit tests.", bold_lead="Testing: ")
bullet("audit_log table exists; confirm privileged actions (ban, role change, grade override, appeal decisions) actually write to it.", bold_lead="Audit trail: ")
bullet("Exam builder does not validate start_time < end_time or that duration fits the window.", bold_lead="Schedule validation: ")
bullet("submitExam has no idempotency guard against double-submit / re-submit after 'submitted'.", bold_lead="Idempotency: ")
bullet("Verify Storage bucket RLS for class_notes and assignment files so students can't read others' uploads.", bold_lead="Storage RLS: ")
bullet("Question bank has a version column but no version-history behaviour; editing a used question can change a past exam's meaning.", bold_lead="Question versioning: ")
hrule()

# 9. BEST FLOWS
doc.add_heading("9. Recommended Best Flows", level=1)
doc.add_heading("9.1 Server-authoritative exam runtime", level=2)
para("Make the server the source of truth for both time and content during an exam:")
numbered("On start, persist started_at and compute deadline = min(started_at + duration, exam.end_time).")
numbered("Send questions WITHOUT answer keys; render from a sanitised payload.")
numbered("Client timer is display-only, seeded from server remaining-time; it re-syncs on refresh.")
numbered("On submit (manual, timeout, or 3-flag auto-submit), the server validates the deadline and grades - late submissions are handled by policy, not trust.")

doc.add_heading("9.2 Exam status lifecycle", level=2)
para("Tighten the state machine and drive transitions by schedule, not manual edits alone:")
code("draft --publish--> upcoming --(start_time reached)--> live\n  live --(end_time reached / all submitted)--> closed\n  closed --(all essays graded)--> graded")
para("A scheduled job (Supabase cron / edge function) can flip upcoming->live->closed so status never lies.")

doc.add_heading("9.3 Submission state machine", level=2)
code("in-progress --submit--> submitted --grade essays--> graded\n  in-progress --3 flags--> flagged\n  flagged --integrity appeal approved--> retake-approved --start--> in-progress")
para("Document these as the only legal transitions and guard each write against illegal jumps.")

doc.add_heading("9.4 Appeal flow (already close to ideal)", level=2)
para("Keep the 7-day window, integrity-before-score ordering, and notifications. Add: surface the deadline "
     "countdown to students, and log every decision to audit_log.")
hrule()

# 10. ROADMAP
doc.add_heading("10. Improvement Roadmap (Prioritised)", level=1)
table(
    ["Priority", "Item", "Effort", "Impact"],
    [
        ["P0", "Strip answer keys from exam payload", "Low", "Critical"],
        ["P0", "Server-enforced exam timing (started_at + deadline)", "Medium", "Critical"],
        ["P0", "Server-side flag/heartbeat; define proctoring as deterrence", "Medium", "High"],
        ["P1", "Regenerate DB types; remove 'as any' casts", "Low", "High"],
        ["P1", "In-handler ownership checks (defence in depth)", "Medium", "High"],
        ["P1", "Remove getTestProfiles debug endpoint", "Low", "Medium"],
        ["P1", "Submit idempotency guard", "Low", "Medium"],
        ["P1", "Exam schedule validation (start<end, duration fits)", "Low", "Medium"],
        ["P2", "Unit tests: grading, appeal window, flag threshold", "Medium", "High"],
        ["P2", "Wire audit_log for all privileged actions", "Medium", "Medium"],
        ["P2", "Realtime live monitor (replace polling)", "Medium", "Medium"],
        ["P2", "Question version history / lock used questions", "Medium", "Medium"],
        ["P2", "Verify Storage bucket RLS", "Low", "High"],
        ["P2", "Email notifications + accessibility pass", "Medium", "Medium"],
    ],
    widths=[0.8, 3.6, 0.9, 1.0],
    header_fill="C02A6B",
)
hrule()

# 11. SUGGESTIONS
doc.add_heading("11. Additional Suggestions", level=1)
bullet("Enrolment by class code/invite so students self-join instead of manual seeding.", bold_lead="UX: ")
bullet("Auto-save in-progress answers (per question) so a crash or refresh doesn't lose work.", bold_lead="Resilience: ")
bullet("Question randomisation / per-student shuffle of options to blunt copying.", bold_lead="Integrity: ")
bullet("A 'grace period' setting is already implied in admin settings - wire it into the timer policy.", bold_lead="Config: ")
bullet("Analytics for lecturers: item difficulty, score distribution, question discrimination.", bold_lead="Insight: ")
bullet("Centralise error handling - many handlers swallow notification errors with empty catch blocks; log them.", bold_lead="Observability: ")
bullet("Add rate-limiting on recordFlag and submit endpoints.", bold_lead="Abuse: ")
bullet("Consider PDF export (not just print) for the score slips using the existing report HTML.", bold_lead="Reporting: ")
hrule()

# 12. CONCLUSION
doc.add_heading("12. Conclusion", level=1)
para(
    "Aura Exam is an impressively complete, well-structured platform with a clean server-function "
    "architecture and a thoughtful, privacy-respecting design. Its grading, appeals, and notification "
    "systems are production-shaped. The gap between 'feature-complete' and 'exam-ready' is concentrated in "
    "three areas: exam-time content security (don't ship answers to the client), server-enforced timing, "
    "and honest framing of browser-based proctoring as deterrence rather than prevention. Closing the P0 "
    "items removes the cheating vectors; the P1 items restore type safety and defence-in-depth; the P2 "
    "items mature it toward a reliable, observable product."
)
para("Addressed in priority order, this system is well within reach of being genuinely exam-ready.", bold=True, color=VIOLET)
doc.add_paragraph().paragraph_format.space_after = Pt(12)
para("- End of report -", align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, italic=True)

out = r"C:\Users\Raziq Shafiee\Desktop\Exam\Aura-Exam-System-Analysis.docx"
doc.save(out)
print("SAVED:", out)
